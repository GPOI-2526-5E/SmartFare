# -*- coding: utf-8 -*-
"""Genera Tesina_SmartFare.docx (~50 pagine) per Esame di Stato."""

from __future__ import annotations

import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_PATH = os.path.join(os.path.dirname(__file__), "Tesina_SmartFare_50_pagine.docx")
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def add_centered(doc: Document, text: str, *, size: int = 11, bold: bool = False, space_after: int = 6) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    p.paragraph_format.space_after = Pt(space_after)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)


def add_image_placeholder(doc: Document, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[ IMMAGINE: {caption} ]")
    r.italic = True
    r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(14)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def read_project_file(rel_path: str, max_lines: int = 0) -> str:
    path = os.path.join(PROJECT_ROOT, rel_path.replace("/", os.sep))
    if not os.path.isfile(path):
        return f"// File non trovato: {rel_path}"
    with open(path, encoding="utf-8", errors="replace") as f:
        lines = f.readlines()
    if max_lines > 0:
        lines = lines[:max_lines]
    return "".join(lines).rstrip()


def add_source_listing(doc: Document, title: str, rel_path: str, explanation: str, max_lines: int = 0) -> None:
    add_heading(doc, title, 2)
    add_body(doc, explanation)
    add_code(doc, f"// Percorso: {rel_path}\n{read_project_file(rel_path, max_lines)}")


def add_long_theory(doc: Document, title: str, paragraphs: list[str]) -> None:
    add_heading(doc, title, 2)
    for p in paragraphs:
        add_body(doc, p)


def add_dispense_flow(doc: Document, title: str, intro: str, steps: list[str]) -> None:
    add_heading(doc, title, 2)
    add_body(doc, intro)
    for i, step in enumerate(steps, 1):
        add_body(doc, f"Passo {i}. {step}")


def build_cover(doc: Document) -> None:
    add_centered(doc, "[Nome Istituto di Istruzione Superiore]", size=12, bold=True, space_after=4)
    add_centered(doc, "[Classe, es. 5°B Informatica]", size=11, space_after=24)
    add_centered(doc, "SmartFare", size=28, bold=True, space_after=8)
    add_centered(
        doc,
        "Piattaforma web AI-driven per la pianificazione\ndi viaggi in Italia",
        size=14,
        space_after=20,
    )
    add_image_placeholder(doc, "Logo SmartFare e schermata Home con barra prompt AI")
    add_centered(doc, "Autori:  [Nome Cognome] — [Eventuale secondo autore]", size=11, space_after=4)
    add_centered(doc, "Anno scolastico:  [2025/26]", size=11, space_after=4)
    add_centered(doc, "Documento:  Tesina Esame di Stato", size=11, bold=True)
    add_page_break(doc)


def build_sommario(doc: Document) -> None:
    add_heading(doc, "Sommario", 1)
    entries = [
        "1. Prefazione e introduzione",
        "2. Perché abbiamo scelto questo progetto",
        "3. Tecnologie utilizzate",
        "   3.1 TypeScript",
        "   3.2 Angular — framework frontend",
        "   3.3 Node.js ed Express — backend",
        "   3.4 PostgreSQL e Prisma ORM",
        "   3.5 Autenticazione: JWT, bcrypt e sessioni",
        "   3.6 Zod — validazione degli input",
        "   3.7 Google Gemini — intelligenza artificiale",
        "   3.8 Leaflet e OpenStreetMap",
        "   3.9 Servizi esterni (Cloudinary, email, OAuth)",
        "4. Architettura del sistema e difficoltà affrontate",
        "5. Il database relazionale",
        "6. Descrizione delle funzionalità",
        "7. Esempi di utilizzo end-to-end",
        "8. Sicurezza, test e considerazioni operative",
        "9. Conclusioni e sviluppi futuri",
        "Allegato A — Struttura del progetto",
        "Allegato B — Checklist screenshot",
        "Allegato C — Webografia",
    ]
    for e in entries:
        p = doc.add_paragraph(e)
        p.paragraph_format.space_after = Pt(2)
    add_page_break(doc)


def section_1(doc: Document) -> None:
    add_heading(doc, "1. Prefazione e introduzione", 1)
    add_body(
        doc,
        "SmartFare è una piattaforma web full-stack per la pianificazione di viaggi in Italia. "
        "L'utente descrive destinazione, date, stile del viaggio e compagno di viaggio; il sistema "
        "combina un database relazionale di località, hotel e attività con modelli di intelligenza "
        "artificiale generativa (Google Gemini) per produrre itinerari strutturati, modificabili "
        "e visualizzabili su mappa.",
    )
    add_body(
        doc,
        "A differenza di un semplice chatbot, SmartFare mantiene uno stato persistente: account utente, "
        "preferenze di viaggio, bozze di itinerari, sessioni di chat e pubblicazione verso la community "
        "(Discover). Il progetto integra competenze del percorso di Informatica: programmazione web "
        "client-server, basi di dati, sicurezza applicativa, API REST, UX e integrazione di servizi cloud.",
    )
    add_body(
        doc,
        "Il documento segue la struttura tipica di una tesina d'esame di stato: motivazioni della scelta, "
        "tecnologie con spiegazione teorica e legame al progetto, architettura software, modello dati, "
        "funzionalità descritte con esempi, difficoltà incontrate e conclusioni. Il riferimento di forma "
        "è l'esempio «Blocco C» (Cairone/Giacchello), adattato a un'applicazione web moderna.",
    )
    add_image_placeholder(doc, "Home page di SmartFare con prompt AI nella hero section")


def section_2(doc: Document) -> None:
    add_heading(doc, "2. Perché abbiamo scelto questo progetto", 1)
    add_body(
        doc,
        "La scelta nasce dall'interesse per il turismo digitale e per l'IA generativa applicata a "
        "problemi concreti — non come assistente generico, ma come strumento che opera su dati reali "
        "(coordinate, categorie, preferenze utente).",
    )
    add_body(doc, "Le motivazioni principali sono quattro:")
    add_bullet(doc, "Problema reale: organizzare un viaggio richiede incrociare date, luoghi, preferenze e offerte locali.")
    add_bullet(doc, "Competenze del percorso: integra frontend, backend, database, sicurezza e API esterne.")
    add_bullet(doc, "Valore didattico: mostra un flusso completo dall'interfaccia alla persistenza su PostgreSQL.")
    add_bullet(
        doc,
        "Differenziazione: l'IA usa POI e hotel già nel database, riducendo allucinazioni e rendendo il risultato mappabile.",
    )
    add_body(
        doc,
        "Nota sull'evoluzione: nelle fasi iniziali il dominio includeva il confronto biglietti ferroviari; "
        "il modello dati è stato riallineato verso itinerari turistici (POI, tappe, visibilità pubblica/privata), "
        "mantenendo l'architettura a tre livelli (Angular, Express, PostgreSQL/Supabase).",
    )


def section_3(doc: Document) -> None:
    add_heading(doc, "3. Tecnologie utilizzate", 1)
    add_body(
        doc,
        "Per ogni tecnologia: definizione, funzionamento generale e motivazione della scelta in SmartFare.",
    )

    tech_sections = [
        (
            "3.1 TypeScript",
            "TypeScript è un superset di JavaScript con tipi statici verificati in compilazione. "
            "Riduce errori su oggetti complessi (Itinerary, ItineraryItem, UserProfile) e migliora "
            "il refactoring in un codebase condiviso tra frontend Angular 21 e backend Express 5.",
        ),
        (
            "3.2 Angular — framework frontend",
            "Angular organizza l'app in componenti standalone, servizi injectable e routing con lazy loading. "
            "I signal e le computed property gestiscono stato reattivo (Discover, builder). "
            "L'HttpClient, con interceptor JWT, comunica con le API sotto /api.",
        ),
        (
            "3.3 Node.js ed Express — backend",
            "Node.js esegue JavaScript/TypeScript lato server. Express 5 espone route modulari, middleware "
            "(helmet, CORS, rate limit globale 50 req/15 min, limiter dedicato AI 20 req/min), "
            "validazione body e gestione errori centralizzata.",
        ),
        (
            "3.4 PostgreSQL e Prisma ORM",
            "PostgreSQL è un DBMS relazionale ACID. Prisma definisce lo schema in schema.prisma, genera "
            "il client type-safe e gestisce migrazioni. Su Supabase il database è hostato in cloud con "
            "backup e connessione SSL.",
        ),
        (
            "3.5 Autenticazione: JWT, bcrypt e sessioni",
            "Alla login si crea un JWT (userId, email, sessionId) e un record AuthSession. "
            "Ogni richiesta protetta verifica firma JWT e che la sessione non sia revocata. "
            "Le password locali usano bcrypt; OAuth Google/GitHub integrato nel frontend.",
        ),
        (
            "3.6 Zod — validazione degli input",
            "Zod definisce schemi per body e query (es. aiItineraryGenerateSchema). "
            "In caso di errore il client riceve 400 con messaggio strutturato, evitando dati inconsistenti nel DB.",
        ),
        (
            "3.7 Google Gemini — intelligenza artificiale",
            "Gemini (API @google/generative-ai) genera itinerari iniziali, risponde in chat contestualizzata "
            "sul workspace (attività/hotel della località) e supporta Voyager AI in streaming. "
            "I prompt includono preferenze utente e vincoli JSON di output.",
        ),
        (
            "3.8 Leaflet e OpenStreetMap",
            "Leaflet renderizza tile OSM, marker e polyline per percorsi. Marker cluster sulla mappa "
            "interattiva d'Italia; drag-and-drop ordine tappe nel builder con aggiornamento rotta.",
        ),
        (
            "3.9 Servizi esterni",
            "Cloudinary per upload avatar/copertine; Nodemailer per verifica email e reset password; "
            "script in utils/ per arricchire località e attività da CSV e Overpass.",
        ),
    ]
    for title, body in tech_sections:
        add_heading(doc, title, 2)
        add_body(doc, body)
        add_body(doc, "")  # spacing

    add_code(
        doc,
        """// Smartfare-Frontend — auth.interceptor.ts (estratto)
export const authInterceptor: HttpInterceptorFn = (req, next) => {
  const token = authService.getAccessToken();
  if (!token || !req.url.startsWith(environment.apiUrl)) return next(req);
  return next(req.clone({ setHeaders: { Authorization: `Bearer ${token}` } }));
};""",
    )
    add_code(
        doc,
        """// Smartfare-Backend — app.ts (estratto)
app.use(helmet());
app.use(globalLimiter);
app.use(cors({ origin: allowedOrigins, credentials: true }));
app.use('/api/ai', aiRoutes);
app.use('/api/itineraries', itineraryRoutes);
app.use('/auth', authRoutes);""",
    )


def section_4(doc: Document) -> None:
    add_heading(doc, "4. Architettura del sistema e difficoltà affrontate", 1)
    add_heading(doc, "4.1 Schema logico client–server–database", 2)
    add_body(doc, "Flusso tipico — generazione itinerario da prompt:")
    for step in [
        "Il componente Angular invoca ai-chat.service.generateItinerary(prompt).",
        "POST /api/ai/itinerary/generate con body validato da Zod.",
        "authenticateJWT / optionalAuthenticateJWT verifica token e sessione.",
        "identifyLocation() individua la città; getWorkspaceData() carica POI e hotel.",
        "Gemini restituisce struttura giorni/tappe; salvataggio opzionale su Itinerary + ItineraryItem.",
        "Il frontend apre builder o preview con mappa aggiornata.",
    ]:
        add_bullet(doc, step)

    add_heading(doc, "4.2 Routing frontend e lazy loading", 2)
    add_code(
        doc,
        """// app.routes.ts (estratto)
{ path: 'voyager', loadComponent: () => import('./features/voyager-ai/...'),
  canActivate: [authGuard] },
{ path: 'discover', loadComponent: () => import('./features/discover/...') },
{ path: 'itineraries/builder', loadComponent: () => import('./features/planner/itinerary-builder/...') },
{ path: 'interactive-map', loadComponent: () => import('./features/interactive-map/...') },""",
    )

    add_heading(doc, "4.3 Difficoltà affrontate", 2)
    difficulties = [
        "Sincronizzazione mappa–lista tappe: invalidateSize() di Leaflet su resize e tab mobile.",
        "Chat AI legata a itinerario: creazione ChatSession al primo messaggio e collegamento chatSessionId su Itinerary.",
        "Rate limiting su endpoint costosi (Gemini) per evitare abuso.",
        "CORS e deploy su Vercel/Render: whitelist origini e variabili ambiente separate.",
        "Popolamento dati territoriali: script Node in utils/ con cache Wikimedia/Overpass.",
    ]
    for d in difficulties:
        add_bullet(doc, d)
    add_image_placeholder(doc, "Diagramma architettura tre livelli (client, API, DB)")


def section_5(doc: Document) -> None:
    add_heading(doc, "5. Il database relazionale", 1)
    add_heading(doc, "5.1 Entità principali", 2)
    entities = [
        "User, UserProfile, UserPreference, UserPreferenceInterest — anagrafica e gusti di viaggio.",
        "Location — comuni italiani con lat/lng e immagine cached.",
        "Activity, ActivityCategory, Accommodation — POI e hotel per località.",
        "Itinerary, ItineraryItem, ItineraryItemType, ItineraryVisibility — percorsi multi-giorno.",
        "ItineraryFavorite, Follow — social e preferiti.",
        "ChatSession, ChatMessage — storico conversazioni AI.",
        "AuthSession — sessioni server-side associate al JWT.",
    ]
    for e in entities:
        add_bullet(doc, e)

    add_heading(doc, "5.2 Normalizzazione e vincoli", 2)
    add_body(
        doc,
        "ItineraryItem referenzia opzionalmente Activity o Accommodation tramite FK; "
        "vincolo unico (itineraryId, dayNumber, orderInt) garantisce ordine tappe per giorno. "
        "Follow usa chiave composita (followerId, followingId). Le visibilità sono tabella lookup "
        "(PRIVATE, PUBLIC, …) invece di stringhe libere.",
    )
    add_code(
        doc,
        """model Itinerary {
  id             Int       @id @default(autoincrement())
  name           String
  visibilityCode String    @default("PRIVATE")
  userId         Int?
  locationId     Int?
  chatSessionId  Int?      @unique
  items          ItineraryItem[]
}""",
    )
    add_image_placeholder(doc, "Diagramma ER (Prisma Studio o draw.io)")


def section_6(doc: Document) -> None:
    add_heading(doc, "6. Descrizione delle funzionalità", 1)
    features = [
        (
            "6.1 Home e accesso rapido all'IA",
            "Hero con carousel destinazioni, barra di ricerca naturale e CTA verso Voyager o generazione diretta. "
            "Sezione SMARTFARE FLOW (scroll storytelling) e itinerari in evidenza.",
        ),
        (
            "6.2 Discover — community",
            "Vetrina itinerari pubblici, ricerca per città/utente/percorso, top like, follow creator, "
            "mappa percorso selezionato, sezione mappa interattiva promozionale.",
        ),
        (
            "6.3 Pianificatore manuale e builder",
            "Flusso: scelta località e date → preview timeline+mappa → builder con sidebar POI, "
            "chat laterale per modifiche in linguaggio naturale, salvataggio bozza.",
        ),
        (
            "6.4 Voyager AI",
            "Assistente conversazionale multi-sessione, pin, streaming risposta, contesto planner/assistant.",
        ),
        (
            "6.5 Profilo e libreria itinerari",
            "Avatar, bio, link social, preferenze viaggio, lista bozze/pubblicati, impostazioni sicurezza.",
        ),
        (
            "6.6 Mappa interattiva d'Italia",
            "Cluster marker per categorie, filtri hotel/ristoranti/attività, popup dettaglio POI.",
        ),
        (
            "6.7 Autenticazione",
            "Registrazione con verifica email OTP, login locale, OAuth Google/GitHub, reset password, revoca sessioni.",
        ),
    ]
    for title, body in features:
        add_heading(doc, title, 2)
        add_body(doc, body)
        add_image_placeholder(doc, title.replace("6.", "Schermata — "))


def section_7(doc: Document) -> None:
    add_heading(doc, "7. Esempi di utilizzo end-to-end", 1)

    add_heading(doc, "Esempio A — Itinerario da prompt AI", 2)
    add_body(doc, "L'utente inserisce: «3 giorni a Roma, ritmo moderato, focus cultura».")
    for s in [
        "Backend identifica locationId Roma e carica workspace.",
        "Gemini propone giorni e tappe con orari indicativi.",
        "Utente affina nel builder e salva; opzionalmente pubblica su Discover.",
    ]:
        add_bullet(doc, s)
    add_code(
        doc,
        """// ai-chat.service.ts (estratto)
generateItinerary(prompt: string): Observable<Itinerary | null> {
  return this.http.post<{ itinerary?: Itinerary }>(
    `${environment.apiUrl}/api/ai/itinerary/generate`, { prompt }
  ).pipe(map((r) => r.itinerary || null));
}""",
    )

    add_heading(doc, "Esempio B — Modifica in chat nel builder", 2)
    add_body(
        doc,
        "Con bozza aperta, l'utente scrive: «Aggiungi un museo il secondo giorno». "
        "POST /api/ai/itinerary/chat invia snapshot tappe + messaggio; la risposta aggiorna mappa e timeline.",
    )

    add_heading(doc, "Esempio C — Discover e follow", 2)
    add_body(
        doc,
        "Utente A pubblica itinerario PUBLIC; B lo trova in Discover, segue A e salva nei preferiti.",
    )


def section_8(doc: Document) -> None:
    add_heading(doc, "8. Sicurezza, test e considerazioni operative", 1)
    add_body(
        doc,
        "Helmet imposta header HTTP sicuri. Rate limit protegge login e AI. "
        "Password mai in chiaro (bcrypt). JWT con scadenza e sessionId revocabile. "
        "Upload file filtrati e inviati su Cloudinary. Validazione Zod su tutti gli input critici.",
    )
    add_body(
        doc,
        "Deploy: frontend su Vercel (SPA), backend su Render, database Supabase PostgreSQL. "
        "Variabili .env per GEMINI_API_KEY, JWT_SECRET, DATABASE_URL, CLOUDINARY, SMTP.",
    )
    add_body(
        doc,
        "Test: Jest configurato sul backend; test manuali su flussi auth, generazione AI e mappe. "
        "Miglioramento futuro: test e2e Cypress/Playwright.",
    )


def section_9(doc: Document) -> None:
    add_heading(doc, "9. Conclusioni e sviluppi futuri", 1)
    add_heading(doc, "Risultati raggiunti", 2)
    for r in [
        "SPA Angular 21 con routing lazy, guard e mappe Leaflet.",
        "API Express tipizzate, Prisma, integrazione Gemini contestualizzata.",
        "Database normalizzato 15+ entità, funzioni social e Discover.",
        "Sicurezza JWT+sessioni, OAuth, verifica email.",
    ]:
        add_bullet(doc, r)

    add_heading(doc, "Limiti attuali", 2)
    for l in [
        "Costi API Gemini e dipendenza da connettività.",
        "Dataset POI aggiornato tramite script manuali.",
        "Nessuna integrazione orari apertura/prenotazioni reali.",
    ]:
        add_bullet(doc, l)

    add_heading(doc, "Sviluppi possibili", 2)
    for s in [
        "Export PDF/ICS itinerario.",
        "Routing OSRM tra tappe.",
        "Recensioni e rating su itinerari pubblici.",
        "Notifiche push promemoria tappe.",
        "Integrazione prezzi trasporti (API ferroviarie).",
    ]:
        add_bullet(doc, s)


def allegati(doc: Document) -> None:
    add_page_break(doc)
    add_heading(doc, "Allegato A — Struttura del progetto", 1)
    add_code(
        doc,
        """SmartFare/
├── Smartfare-Frontend/src/app/
│   ├── core/          (services, interceptors, guards, i18n)
│   └── features/
│       ├── home/ discover/ planner/ voyager-ai/
│       ├── profile/ interactive-map/ auth/
└── Smartfare-Backend/
    ├── src/routes/ services/ middleware/ schemas/
    └── prisma/schema.prisma""",
    )

    add_heading(doc, "Allegato B — Checklist screenshot", 1)
    shots = [
        "Home con prompt AI",
        "Login e registrazione",
        "Discover — griglia e mappa percorso",
        "Manual planner e builder con chat",
        "Preview timeline + polyline",
        "Voyager AI — sessioni e streaming",
        "Mappa interattiva Italia — cluster",
        "Profilo e preferenze viaggio",
        "Prisma Studio / diagramma ER",
    ]
    for s in shots:
        add_bullet(doc, f"☐  {s}")

    add_heading(doc, "Allegato C — Webografia", 1)
    links = [
        "Angular — https://angular.dev",
        "Node.js — https://nodejs.org",
        "Prisma — https://www.prisma.io",
        "PostgreSQL — https://www.postgresql.org",
        "Google Gemini API — https://ai.google.dev",
        "Leaflet — https://leafletjs.com",
        "OpenStreetMap — https://www.openstreetmap.org",
        "Supabase — https://supabase.com",
    ]
    for l in links:
        add_bullet(doc, l)


def add_detailed_module_sections(doc: Document) -> None:
    """Sezioni stile «dispense» Cairone — dettaglio moduli software."""
    add_page_break(doc)
    add_heading(doc, "Appendice didattica — Dettaglio moduli software", 1)

    modules = [
        (
            "Modulo Home (home-section)",
            "Componente principale della landing: carousel hero con località italiane, "
            "form di ricerca collegato al flusso AI, sezione features-grid con scroll storytelling "
            "(tre scene: planner, libreria itinerari, Voyager AI), itinerari in evidenza e footer SEO. "
            "Utilizza AOS per animazioni on-scroll e servizi itinerary/location per dati dinamici.",
        ),
        (
            "Modulo Discover (discover-page)",
            "Pagina vetrina post-login o pubblica: hero con autocomplete città, griglia itinerari "
            "più votati (featured-mosaic), pannello mappa percorso selezionato, rail creator seguiti, "
            "suggerimenti «vicino a te» se loggato, ultimi pubblicati. Integra discover-map con Leaflet "
            "e API search/follow/favorites.",
        ),
        (
            "Modulo Planner (manual-planner, itinerary-builder, preview)",
            "manual-planner: scelta Location e intervallo date. "
            "itinerary-builder: workspace con mappa, sidebar POI filtrabili, chat AI laterale, "
            "riordino tappe drag-and-drop. itinerary-preview: timeline giorni e export verso salvataggio. "
            "Tutti condividono itinerary.service e modelli ItineraryWorkspace.",
        ),
        (
            "Modulo Voyager AI (voyager-ai)",
            "Interfaccia chat stile assistente: lista sessioni pin/abilitate, streaming SSE o chunk "
            "dalla API chat, creazione nuova sessione con titolo auto. Modalità planner vs assistant "
            "cambia il system prompt lato backend.",
        ),
        (
            "Modulo Mappa interattiva (interactive-map)",
            "Vista nazionale con cluster marker per categoria, filtri sidebar, caricamento "
            "attività/alloggi per bounding box. Utilizza categoryVisuals per icone e colori coerenti "
            "con il builder.",
        ),
        (
            "Modulo Auth e Profilo",
            "login/register/oauth-callback, verify-email, forgot/reset password. "
            "profile-view con follower, settings con preferenze viaggio e gestione sessioni attive. "
            "upload avatar tramite endpoint /api/upload e Cloudinary.",
        ),
        (
            "Backend — itinerary.service",
            "Logica di business: CRUD itinerari, composizione workspace per AI, calcolo timing tappe "
            "(itinerary-item-timing.util), slug pubblici, visibilità, popolamento item da attività/hotel.",
        ),
        (
            "Backend — gemini.service",
            "Costruzione prompt con blocco preferenze utente, identifyLocation (regex + fallback LLM), "
            "generateInitialItinerary, chat iterativa con JSON schema, gestione errori e retry.",
        ),
    ]
    for title, text in modules:
        add_heading(doc, title, 2)
        add_body(doc, text)
        add_body(
            doc,
            "Dal punto di vista del testing manuale, ogni modulo è stato verificato con casi "
            "nominali (utente autenticato, dati presenti in DB) e casi limite (prompt senza città, "
            "sessione scaduta, località senza attività). Gli errori HTTP sono gestiti da error.middleware "
            "che restituisce JSON { error, statusCode } uniforme al frontend.",
        )
        add_image_placeholder(doc, f"Screenshot — {title}")


def add_extended_source_volume(doc: Document) -> None:
    """Listati sorgente e dispense operative — volume stile tesina Cairone."""
    add_page_break(doc)
    add_heading(doc, "Parte II — Listati di codice commentati", 1)
    add_body(
        doc,
        "Come nell'esempio di tesina «Blocco C», in cui ogni funzione del software è spiegata "
        "affiancando il codice C# di traduzione, qui riportiamo estratti commentati dei file "
        "più rappresentativi del progetto SmartFare. I listati sono tratti dal repository reale.",
    )

    add_source_listing(
        doc,
        "Listato 1 — Middleware di autenticazione JWT",
        "Smartfare-Backend/src/middleware/auth.middleware.ts",
        "Il middleware estrae il Bearer token, verifica la firma JWT con JWT_SECRET, "
        "controlla che la sessione non sia revocata su tabella AuthSession e popola req.user "
        "per i controller successivi. optionalAuthenticateJWT permette route pubbliche con utente opzionale.",
        0,
    )
    add_page_break(doc)
    add_source_listing(
        doc,
        "Listato 2 — Schema Prisma (database)",
        "Smartfare-Backend/prisma/schema.prisma",
        "Definizione dichiarativa di tutte le entità: utenti, preferenze, località, attività, "
        "hotel, itinerari, tappe, chat, follow. Prisma genera il client TypeScript e applica migrazioni.",
        0,
    )
    add_page_break(doc)
    add_source_listing(
        doc,
        "Listato 3 — Routing Angular",
        "Smartfare-Frontend/src/app/app.routes.ts",
        "Ogni route usa loadComponent per code splitting. authGuard protegge voyager e libreria itinerari.",
        0,
    )
    add_page_break(doc)
    add_source_listing(
        doc,
        "Listato 4 — Route AI (generazione e chat)",
        "Smartfare-Backend/src/routes/ai.route.ts",
        "Endpoint /itinerary/generate e /itinerary/chat con rate limit, validazione Zod, "
        "caricamento workspace e integrazione GeminiItineraryChatService.",
        180,
    )
    add_page_break(doc)
    add_source_listing(
        doc,
        "Listato 5 — Servizio Gemini (estratto)",
        "Smartfare-Backend/src/services/ai/gemini.service.ts",
        "Gestione modelli Gemini, retry con backoff, costruzione prompt con preferenze utente "
        "e parsing JSON della risposta itinerario.",
        220,
    )

    add_page_break(doc)
    add_heading(doc, "Parte III — Dispense operative (uso del software)", 1)
    add_body(
        doc,
        "Sezione analogica alla «Dispense Blocco C» dell'esempio Cairone: per ogni area funzionale "
        "descriviamo passo per passo cosa fa l'utente e cosa succede dietro le quinte.",
    )

    add_dispense_flow(
        doc,
        "A. Registrazione e primo accesso",
        "Percorso per un nuovo utente che vuole usare planner e Voyager.",
        [
            "L'utente apre /register e compila email, password e dati profilo base.",
            "Il backend valida con Zod, hasha la password (bcrypt) e crea record User + UserProfile.",
            "Viene inviata email di verifica con token temporaneo (Nodemailer).",
            "L'utente clicca il link in /verify-email; il flag isEmailVerified passa a true.",
            "Al login POST /auth/login crea AuthSession e restituisce JWT al client.",
            "L'interceptor Angular allega Authorization: Bearer a ogni chiamata API.",
            "L'utente viene reindirizzato alla home o al planner secondo il deep link.",
            "In impostazioni può collegare OAuth Google/GitHub per login alternativo.",
            "Può impostare preferenze viaggio (stile, ritmo, compagno) usate nei prompt AI.",
            "Le sessioni attive sono elencabili e revocabili dal profilo sicurezza.",
        ],
    )
    add_dispense_flow(
        doc,
        "B. Creazione itinerario manuale",
        "Percorso senza AI: scelta esplicita di destinazione e date.",
        [
            "Da home o menu si apre /itineraries/new (manual-planner).",
            "Autocomplete cerca Location nel database (nome comune, provincia).",
            "Si scelgono data inizio e fine; il client valida intervallo.",
            "Si passa a preview con timeline vuota e mappa centrata sulla città.",
            "Dal builder si apre sidebar: filtri per categoria attività.",
            "Clic su POI aggiunge ItineraryItem con dayNumber e orderInt.",
            "Drag sulla mappa riordina tappe; il client ricalcola polyline.",
            "Salvataggio POST /api/itineraries persiste itinerario e items.",
            "Si può cambiare visibilità PRIVATE/PUBLIC per Discover.",
            "Anteprima finale mostra riepilogo per giorno prima della pubblicazione.",
        ],
    )
    add_dispense_flow(
        doc,
        "C. Generazione itinerario con intelligenza artificiale",
        "Percorso principale dimostrato all'esame.",
        [
            "L'utente scrive prompt naturale nella home o in Voyager.",
            "Frontend chiama POST /api/ai/itinerary/generate.",
            "Backend identifica città (regex + Gemini) tra le Location.",
            "getWorkspaceData carica attività, hotel e preferenze.",
            "Gemini restituisce struttura JSON giorni/tappe/alloggi.",
            "Risposta include locationId e bozza itinerario.",
            "Redirect al builder con dati precompilati.",
            "Utente modifica tappe manualmente o via chat laterale.",
            "Ogni messaggio chat invia snapshot corrente a /api/ai/itinerary/chat.",
            "Salvataggio definitivo e eventuale pubblicazione in Discover.",
        ],
    )
    add_dispense_flow(
        doc,
        "D. Discover e funzioni social",
        "Esplorazione community.",
        [
            "Apertura /discover senza login (vetrina pubblica).",
            "Hero con carousel e barra ricerca città/itinerari/utenti.",
            "Sezione top like mostra itinerari con più preferiti.",
            "Click su card seleziona itinerario e disegna rotta su mappa.",
            "Pulsante anteprima apre itinerary-preview read-only.",
            "Follow su autore tramite API profile/follow.",
            "Ricerca avanzata con tab itinerari / utenti / luoghi.",
            "Utente loggato vede suggerimenti «vicino ai tuoi itinerari».",
            "Banner invita alla mappa interattiva nazionale.",
            "Creator più seguiti in rail orizzontale.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "Parte IV — Approfondimenti teorici aggiuntivi", 1)

    add_long_theory(
        doc,
        "4.1 Il paradigma Single Page Application",
        [
            "Una SPA carica una shell HTML e successivamente solo i bundle JavaScript necessari "
            "per la route corrente. In Angular 21 il router usa loadComponent per import dinamici: "
            "la prima visita a /discover scarica solo il chunk discover-page, riducendo il tempo "
            "di first contentful paint rispetto a un bundle monolitico.",
            "Il trade-off è la complessità del routing e della gestione stato: SmartFare usa servizi "
            "singleton (auth, itinerary, ai-chat) e signal nei componenti più interattivi. "
            "Il browser mantiene lo stato JWT in memoria/localStorage mentre PostgreSQL mantiene "
            "lo stato di business (itinerari, chat).",
            "Per l'esame è importante saper spiegare che la SPA non significa «tutto client»: "
            "le operazioni sensibili (AI, pagamenti futuri, validazione) restano server-side.",
        ],
    )
    add_long_theory(
        doc,
        "4.2 REST e risorse del dominio viaggio",
        [
            "REST organizza le API attorno a risorse: itineraries, locations, activities. "
            "I verbi HTTP esprimono intenzione: GET lettura, POST creazione, PUT aggiornamento, DELETE rimozione.",
            "SmartFare non è iper-REST puro (es. /api/ai/itinerary/generate è un RPC-style endpoint) "
            "ma mantiene coerenza JSON e codici stato HTTP (400 validazione, 401 auth, 429 rate limit).",
            "La documentazione implicita è negli schema Zod condivisi tra route e test manuali Postman.",
        ],
    )
    add_long_theory(
        doc,
        "4.3 Intelligenza artificiale generativa e grounding",
        [
            "I Large Language Model predicono testo statisticamente plausibile; senza ancoraggio "
            "possono inventare attrazioni inesistenti. SmartFare applica «grounding»: il prompt "
            "include solo attività e hotel presenti nel workspace della località scelta.",
            "Il modello Gemini 2.x riceve istruzioni di output JSON con schema rigido; il backend "
            "valida e persiste solo strutture parseabili. In caso di quota API, gemini.service "
            "implementa retry con RetryInfo e fallback tra modelli flash.",
            "Questo equilibrio tra creatività dell'utente (linguaggio naturale) e vincoli del DB "
            "è il cuore didattico del progetto per l'esame di stato informatica.",
        ],
    )


def add_technology_deep_dive(doc: Document) -> None:
    add_page_break(doc)
    add_heading(doc, "3 bis — Approfondimento tecnologie (esteso)", 1)
    dives = {
        "TypeScript e tipizzazione del dominio": [
            "I tipi permettono di modellare Itinerary, ItineraryItem, UserProfile come interfacce "
            "condivise. In fase di build il compilatore segnala errori (es. campo mancante su item).",
            "Nel backend i DTO Zod inferiscono tipi da schema, allineando validazione runtime e compile-time.",
        ],
        "Angular: componenti e dependency injection": [
            "Ogni feature è un componente standalone che dichiara imports espliciti (CommonModule, RouterLink).",
            "I servizi providedIn root sono singleton: AuthService mantiene token, ItineraryService cache bozze.",
            "I guard implementano CanActivateFn per bloccare route senza login.",
        ],
        "Express middleware chain": [
            "Ogni richiesta attraversa helmet, rateLimit, cors, json parser, poi router specifico.",
            "Gli errori thrown come AppError vengono catturati da error.middleware con status coerente.",
        ],
        "PostgreSQL e indici": [
            "Indici su userId, itineraryId, followerId accelerano query Discover e profilo.",
            "Le FK con onDelete Cascade evitano orphan record quando si elimina un utente.",
        ],
    }
    for title, paras in dives.items():
        add_long_theory(doc, title, paras)


def add_entity_catalog(doc: Document) -> None:
    add_page_break(doc)
    add_heading(doc, "Appendice D — Dizionario dati (entità e campi)", 1)
    add_body(
        doc,
        "Descrizione campo per campo delle principali tabelle, utile per l'esame orale sul database.",
    )
    entities = [
        ("User", [
            ("id", "Chiave primaria autoincrementale."),
            ("email", "Identificativo univoco per login."),
            ("passwordHash", "Hash bcrypt; null se solo OAuth."),
            ("authProvider", "local | google | github."),
            ("isEmailVerified", "Gate per funzioni sensibili."),
            ("resetPasswordToken / Expires", "OTP reset password temporaneo."),
        ]),
        ("UserProfile", [
            ("name, surname", "Nome visualizzato in Discover e itinerari."),
            ("avatarUrl", "URL Cloudinary."),
            ("bio, instagramUrl, twitterUrl", "Social e presentazione."),
            ("birthDate", "Opzionale per personalizzazione."),
        ]),
        ("UserPreference", [
            ("travelStyle", "Culturale, relax, avventura…"),
            ("pace", "Lento, moderato, intenso."),
            ("travelCompanion", "Solo, coppia, famiglia…"),
            ("UserPreferenceInterest", "Tabella ponte verso ActivityCategory."),
        ]),
        ("Location", [
            ("name, province, cap", "Identificazione comune italiano."),
            ("latitude, longitude", "Centroide per mappa e query spaziali."),
            ("image", "Copertina cached (es. Unsplash)."),
        ]),
        ("Activity / Accommodation", [
            ("name, street", "Denominazione e indirizzo POI."),
            ("latitude, longitude", "Posizione marker Leaflet."),
            ("categoryId", "FK verso ActivityCategory."),
            ("imageUrl", "Media di default o specifico."),
        ]),
        ("Itinerary", [
            ("name, description", "Titolo e testo pubblico."),
            ("visibilityCode", "PRIVATE | PUBLIC (lookup table)."),
            ("publicSlug", "URL amichevole per condividere."),
            ("startDate, endDate", "Intervallo viaggio."),
            ("chatSessionId", "Collegamento chat AI dell'itinerario."),
            ("userId, locationId", "Proprietario e destinazione principale."),
        ]),
        ("ItineraryItem", [
            ("dayNumber, orderInt", "Giorno e ordine tappa."),
            ("itemTypeCode", "ACTIVITY | ACCOMMODATION | …"),
            ("activityId / accommodationId", "FK opzionali verso POI."),
            ("plannedStartAt, plannedEndAt", "Orari pianificati."),
            ("groupName", "Raggruppamento tappe (es. mattina)."),
        ]),
        ("ChatSession / ChatMessage", [
            ("mode", "planner | assistant | itinerary."),
            ("locationId", "Contesto geografico."),
            ("role", "user | assistant sui messaggi."),
            ("isPinned", "Sessioni Voyager in evidenza."),
        ]),
        ("Follow / ItineraryFavorite", [
            ("followerId, followingId", "Relazione social tra utenti."),
            ("userId, itineraryId", "Preferiti su itinerari altrui."),
        ]),
        ("AuthSession", [
            ("id UUID", "Identificatore sessione nel JWT."),
            ("revokedAt", "Logout forzato o revoca da impostazioni."),
            ("userAgent, ipAddress", "Audit sicurezza."),
        ]),
    ]
    for entity, fields in entities:
        add_heading(doc, entity, 2)
        for fname, fdesc in fields:
            add_bullet(doc, f"{fname}: {fdesc}")


def add_oral_exam_prep(doc: Document) -> None:
    add_page_break(doc)
    add_heading(doc, "Appendice E — Domande frequenti all'esame orale", 1)
    qa = [
        (
            "Perché avete scelto Angular e non React?",
            "Angular offre out-of-the-box routing, form, HTTP e struttura modulare adatta a un team scolastico "
            "che deve consegnare un prodotto completo. React è più flessibile ma richiede più scelte architetturali. "
            "Per SmartFare la coerenza del framework e TypeScript nativo hanno accelerato lo sviluppo di mappe, "
            "form complessi e guard di autenticazione.",
        ),
        (
            "Come evitate che l'IA inventi luoghi inesistenti?",
            "Il backend non chiede a Gemini di inventare POI: gli passa un workspace con elenco attività e hotel "
            "già presenti nel database per quella Location. Il modello seleziona e combina elementi esistenti. "
            "Se la città non è nel DB, identifyLocation fallisce o chiede chiarimento.",
        ),
        (
            "Spiegate il flusso JWT + sessione server-side.",
            "Il JWT contiene userId, email e sessionId. Oltre alla firma crittografica, il middleware verifica "
            "che sessionId esista in AuthSession e non sia revocato. Così un token rubato può essere invalidato "
            "dal server senza attendere la scadenza naturale del JWT.",
        ),
        (
            "Quali misure di sicurezza avete implementato?",
            "Helmet per header HTTP, rate limiting globale e su AI, bcrypt sulle password, CORS ristretto, "
            "validazione Zod, limit upload, variabili segrete in .env, HTTPS in produzione.",
        ),
        (
            "Come funziona la mappa nel builder?",
            "Leaflet carica tile OpenStreetMap. I marker rappresentano ItineraryItem; la polyline collega "
            "le coordinate in ordine di orderInt. Al drag si aggiorna l'ordine e si richiama invalidateSize "
            "per ridisegnare correttamente dopo resize pannello.",
        ),
        (
            "Cosa fareste con più tempo?",
            "Export PDF/ICS, routing OSRM, recensioni itinerari, test automatici e2e, PWA offline, "
            "integrazione orari apertura da API esterne.",
        ),
    ]
    for q, a in qa:
        add_heading(doc, q, 2)
        add_body(doc, a)


def add_readme_embed(doc: Document) -> None:
    add_page_break(doc)
    add_heading(doc, "Appendice F — Documentazione di progetto (README)", 1)
    add_body(doc, "Testo integrale del README del repository SmartFare:")
    add_code(doc, read_project_file("README.md", 0))


def add_api_reference(doc: Document) -> None:
    add_page_break(doc)
    add_heading(doc, "Appendice — Principali endpoint API REST", 1)
    add_body(
        doc,
        "Il backend espone risorse sotto /api e autenticazione sotto /auth. "
        "Di seguito una tabella riassuntiva (metodo — percorso — scopo).",
    )
    endpoints = [
        ("POST", "/auth/register", "Registrazione utente locale"),
        ("POST", "/auth/login", "Login e rilascio JWT + sessione"),
        ("POST", "/auth/oauth/google", "Login social Google"),
        ("GET", "/api/itineraries", "Lista itinerari utente"),
        ("POST", "/api/itineraries", "Crea itinerario"),
        ("PUT", "/api/itineraries/:id", "Aggiorna itinerario e tappe"),
        ("POST", "/api/ai/itinerary/generate", "Genera itinerario da prompt"),
        ("POST", "/api/ai/itinerary/chat", "Modifica itinerario via chat"),
        ("GET", "/api/chat/sessions", "Sessioni Voyager utente"),
        ("GET", "/api/locations", "Elenco località / ricerca"),
        ("GET", "/api/activity", "Attività per località"),
        ("GET", "/api/accommodation", "Hotel per località"),
        ("POST", "/api/profile/follow/:id", "Segui utente"),
        ("POST", "/api/upload", "Upload immagine profilo"),
    ]
    for method, path, desc in endpoints:
        add_bullet(doc, f"{method} {path} — {desc}")


def expand_for_length(doc: Document) -> None:
    """Paragrafi aggiuntivi per avvicinarsi a ~50 pagine."""
    add_page_break(doc)
    add_heading(doc, "Approfondimento — User experience e design", 1)
    add_body(
        doc,
        "L'interfaccia usa tema scuro, gradienti viola/teal e componenti glassmorphism coerenti "
        "con il brand SmartFare. Bootstrap Icons e font Outfit/Playfair Display distinguono "
        "titoli editoriali (itinerari) da UI funzionale. Le animazioni AOS sulla home guidano "
        "l'attenzione senza compromettere le performance (lazy load immagini, marker cluster).",
    )
    add_body(
        doc,
        "Su mobile il builder passa a tab (mappa / strumenti / chat) per evitare sovrapposizioni; "
        "la navbar collassabile e il footer con link SEO completano l'esperienza. "
        "L'i18n (servizio traduzioni) prepara l'app a messaggi multilingua anche se la v1 è prevalentemente italiana.",
    )

    add_heading(doc, "Approfondimento — Pipeline dati territoriali", 1)
    add_body(
        doc,
        "La cartella utils/ contiene script per importare Location da CSV, arricchire attività "
        "via Overpass/Wikidata e verificare output (check_output.mjs). Questo separa "
        "la manutenzione del dataset dall'applicazione runtime: il team può aggiornare POI "
        "senza ridistribuire il frontend.",
    )
    add_body(
        doc,
        "Ogni Activity è legata a ActivityCategory (es. cultura, gastronomia) e a Location; "
        "le preferenze utente (UserPreferenceInterest) filtrano quali categorie privilegiare "
        "nei prompt Gemini, personalizzando i suggerimenti.",
    )

    add_heading(doc, "Approfondimento — Confronto con l'esempio tesina Cairone", 1)
    add_body(
        doc,
        "L'esempio «Blocco C» (43 pagine, Vallauri Fossano 2015/16) documenta un software desktop C#/C "
        "per insegnare ANSI C tramite blocchi visuali. Struttura: copertina, sommario, prefazione, "
        "motivazioni, tecnologie (C, C#, GCC, XML), architettura con difficoltà e codice, "
        "poi sezione estesa «Dispense» con ogni funzione del prodotto.",
    )
    add_body(
        doc,
        "SmartFare replica quella struttura adattata al web: al posto dei blocchi ANSI ci sono "
        "moduli Angular; al posto della conversione XML→C c'è la pipeline prompt→JSON itinerario; "
        "la sezione 6 equivale alle «dispense» con Home, Discover, Builder, Voyager, Mappa. "
        "Il volume è portato a ~50 pagine con approfondimenti su UX, dati, sicurezza e allegati.",
    )

    add_heading(doc, "Glossario", 1)
    glossary = [
        ("POI", "Point Of Interest — punto di interesse turistico."),
        ("SPA", "Single Page Application — applicazione web a pagina unica."),
        ("JWT", "JSON Web Token — token firmato per autenticazione stateless."),
        ("ORM", "Object-Relational Mapping — Prisma mappa tabelle a oggetti TypeScript."),
        ("REST", "Architettura API basata su risorse HTTP (GET, POST, …)."),
    ]
    for term, defn in glossary:
        p = doc.add_paragraph()
        r = p.add_run(f"{term}: ")
        r.bold = True
        p.add_run(defn)


def main() -> None:
    doc = Document()
    set_doc_defaults(doc)

    build_cover(doc)
    build_sommario(doc)
    section_1(doc)
    section_2(doc)
    add_page_break(doc)
    section_3(doc)
    add_technology_deep_dive(doc)
    add_page_break(doc)
    section_4(doc)
    add_extended_source_volume(doc)
    section_5(doc)
    add_page_break(doc)
    section_6(doc)
    add_page_break(doc)
    section_7(doc)
    section_8(doc)
    section_9(doc)
    add_detailed_module_sections(doc)
    add_api_reference(doc)
    add_entity_catalog(doc)
    add_oral_exam_prep(doc)
    add_readme_embed(doc)
    expand_for_length(doc)
    allegati(doc)

    # Footer metadati
    doc.core_properties.title = "SmartFare — Tesina Esame di Stato"
    doc.core_properties.subject = "Piattaforma web AI-driven per viaggi in Italia"
    doc.core_properties.author = "[Nome Cognome]"
    doc.core_properties.comments = f"Generato il {datetime.now().strftime('%d/%m/%Y %H:%M')}"

    doc.save(OUT_PATH)
    print(f"Salvato: {OUT_PATH}")
    print(f"Paragrafi totali: {len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
